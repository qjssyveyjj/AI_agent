import io


class UnsupportedFileType(Exception):
    pass


def extract_text(filename: str, data: bytes) -> str:
    """按扩展名抽取纯文本。"""
    name = filename.lower()
    if name.endswith(".pdf"):
        return _from_pdf(data)
    if name.endswith(".docx"):
        return _from_docx(data)
    if name.endswith(".xlsx"):
        return _from_xlsx(data)
    if name.endswith((".txt", ".md", ".csv", ".json", ".log")):
        return data.decode("utf-8", errors="ignore")
    raise UnsupportedFileType(f"不支持的文件类型：{filename}")


def _from_pdf(data: bytes) -> str:
    from pypdf import PdfReader

    reader = PdfReader(io.BytesIO(data))
    parts = [page.extract_text() or "" for page in reader.pages]
    return "\n".join(parts)


def _from_docx(data: bytes) -> str:
    import docx

    doc = docx.Document(io.BytesIO(data))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_xlsx(data: bytes) -> str:
    from openpyxl import load_workbook

    wb = load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts: list[str] = []
    for ws in wb.worksheets:
        parts.append(f"# 工作表: {ws.title}")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) for c in row if c is not None]
            if cells:
                parts.append(" | ".join(cells))
    wb.close()
    return "\n".join(parts)
